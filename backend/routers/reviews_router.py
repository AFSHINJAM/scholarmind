from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from pydantic import BaseModel
from typing import Optional
from datetime import datetime
import io
from models.database import get_db, Review, Paper
from auth import get_current_user

router = APIRouter()

class ReviewUpdate(BaseModel):
    summary: Optional[str] = None
    major_concerns: Optional[str] = None
    minor_concerns: Optional[str] = None
    recommendation: Optional[str] = None
    score_novelty: Optional[float] = None
    score_methodology: Optional[float] = None
    score_clarity: Optional[float] = None
    score_statistics: Optional[float] = None

@router.get("/")
def get_reviews(db: Session = Depends(get_db), current_user = Depends(get_current_user)):
    reviews = db.query(Review).filter(Review.reviewer_id == current_user.id).all()
    return [serialize_review(r) for r in reviews]

@router.get("/{review_id}")
def get_review(review_id: int, db: Session = Depends(get_db), current_user = Depends(get_current_user)):
    review = db.query(Review).filter(Review.id == review_id, Review.reviewer_id == current_user.id).first()
    if not review: raise HTTPException(404, "Review not found")
    return serialize_review(review)

@router.put("/{review_id}")
def update_review(review_id: int, body: ReviewUpdate, db: Session = Depends(get_db), current_user = Depends(get_current_user)):
    review = db.query(Review).filter(Review.id == review_id, Review.reviewer_id == current_user.id).first()
    if not review: raise HTTPException(404, "Review not found")
    for field, value in body.dict(exclude_none=True).items():
        setattr(review, field, value)
    review.updated_at = datetime.utcnow()
    db.commit()
    return serialize_review(review)

@router.post("/{review_id}/complete")
def complete_review(review_id: int, db: Session = Depends(get_db), current_user = Depends(get_current_user)):
    review = db.query(Review).filter(Review.id == review_id, Review.reviewer_id == current_user.id).first()
    if not review: raise HTTPException(404)
    review.status = "completed"
    if review.paper: review.paper.status = "completed"
    db.commit()
    return {"ok": True}

@router.get("/{review_id}/export")
def export_review(
    review_id: int,
    format: str = Query("docx", enum=["docx", "pdf"]),
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    review = db.query(Review).filter(Review.id == review_id, Review.reviewer_id == current_user.id).first()
    if not review: raise HTTPException(404)

    paper_title = review.paper.title if review.paper else "Unknown"
    journal = review.paper.journal if review.paper else "Unknown"

    if format == "docx":
        return export_docx(review, paper_title, journal, current_user.name)
    else:
        return export_pdf_review(review, paper_title, journal, current_user.name)

def export_docx(review, paper_title, journal, reviewer_name):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Page margins
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

        # Title
        title = doc.add_heading('PEER REVIEW REPORT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata table
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        meta = [
            ('Manuscript', paper_title),
            ('Journal', journal),
            ('Review Date', datetime.utcnow().strftime('%B %d, %Y')),
            ('Reviewer', reviewer_name),
            ('Recommendation', review.recommendation or 'Not specified'),
        ]
        for i, (label, value) in enumerate(meta):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value
        doc.add_paragraph()

        # Sections
        sections = [
            ('SUMMARY', review.summary),
            ('MAJOR CONCERNS', review.major_concerns),
            ('MINOR CONCERNS', review.minor_concerns),
        ]
        for heading, content in sections:
            if content and content.strip():
                doc.add_heading(heading, level=1)
                doc.add_paragraph(content)
                doc.add_paragraph()

        # Scores
        doc.add_heading('SCORES', level=1)
        score_table = doc.add_table(rows=1, cols=2)
        score_table.style = 'Table Grid'
        hdr = score_table.rows[0].cells
        hdr[0].text = 'Criterion'
        hdr[1].text = 'Score (0-10)'
        scores = [
            ('Novelty & Originality', review.score_novelty),
            ('Methodology', review.score_methodology),
            ('Clarity & Presentation', review.score_clarity),
            ('Statistical Rigor', review.score_statistics),
        ]
        for label, score in scores:
            row = score_table.add_row().cells
            row[0].text = label
            row[1].text = str(score) if score is not None else 'N/A'

        doc.add_paragraph()
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_para.add_run('Generated by ScholarMind · Powered by Claude AI · Anthropic')
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(150, 150, 150)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        filename = f"review_{paper_title[:30].replace(' ','_')}.docx"
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        raise HTTPException(500, f"Export failed: {str(e)}")

def export_pdf_review(review, paper_title, journal, reviewer_name):
    try:
        # Generate HTML then convert
        html = f"""
        <html><head><style>
        body{{font-family:Georgia,serif;margin:60px;color:#1a1a1a;line-height:1.6}}
        h1{{text-align:center;font-size:22px;border-bottom:2px solid #1a4a4a;padding-bottom:10px;color:#1a4a4a}}
        h2{{font-size:16px;color:#1a4a4a;margin-top:24px;border-bottom:1px solid #ddd;padding-bottom:6px}}
        table{{width:100%;border-collapse:collapse;margin:16px 0}}
        td,th{{border:1px solid #ddd;padding:8px 12px;font-size:13px}}
        th{{background:#f0f0f0;font-weight:600}}
        p{{font-size:13px;margin:8px 0}}
        .footer{{text-align:center;color:#999;font-size:11px;margin-top:40px;border-top:1px solid #eee;padding-top:12px}}
        </style></head><body>
        <h1>PEER REVIEW REPORT</h1>
        <table>
          <tr><td><strong>Manuscript</strong></td><td>{paper_title}</td></tr>
          <tr><td><strong>Journal</strong></td><td>{journal}</td></tr>
          <tr><td><strong>Review Date</strong></td><td>{datetime.utcnow().strftime('%B %d, %Y')}</td></tr>
          <tr><td><strong>Reviewer</strong></td><td>{reviewer_name}</td></tr>
          <tr><td><strong>Recommendation</strong></td><td><strong>{review.recommendation or 'Not specified'}</strong></td></tr>
        </table>
        <h2>SUMMARY</h2><p>{(review.summary or '').replace(chr(10),'<br>')}</p>
        <h2>MAJOR CONCERNS</h2><p>{(review.major_concerns or '').replace(chr(10),'<br>')}</p>
        <h2>MINOR CONCERNS</h2><p>{(review.minor_concerns or '').replace(chr(10),'<br>')}</p>
        <h2>SCORES</h2>
        <table>
          <tr><th>Criterion</th><th>Score (0-10)</th></tr>
          <tr><td>Novelty & Originality</td><td>{review.score_novelty or 'N/A'}</td></tr>
          <tr><td>Methodology</td><td>{review.score_methodology or 'N/A'}</td></tr>
          <tr><td>Clarity & Presentation</td><td>{review.score_clarity or 'N/A'}</td></tr>
          <tr><td>Statistical Rigor</td><td>{review.score_statistics or 'N/A'}</td></tr>
        </table>
        <div class="footer">Generated by ScholarMind · Powered by Claude AI · Anthropic</div>
        </body></html>"""

        # Try to use weasyprint, fall back to returning HTML as download
        try:
            import weasyprint
            pdf_bytes = weasyprint.HTML(string=html).write_pdf()
            buf = io.BytesIO(pdf_bytes)
            filename = f"review_{paper_title[:30].replace(' ','_')}.pdf"
            return StreamingResponse(buf, media_type="application/pdf",
                headers={"Content-Disposition": f"attachment; filename={filename}"})
        except ImportError:
            # weasyprint not available, send as HTML file
            buf = io.BytesIO(html.encode('utf-8'))
            return StreamingResponse(buf, media_type="text/html",
                headers={"Content-Disposition": f"attachment; filename=review.html"})
    except Exception as e:
        raise HTTPException(500, f"PDF export failed: {str(e)}")

def serialize_review(r):
    return {
        "id": r.id, "paper_id": r.paper_id,
        "paper_title": r.paper.title if r.paper else "",
        "paper_journal": r.paper.journal if r.paper else "",
        "paper_type": r.paper.paper_type if r.paper else "",
        "summary": r.summary, "major_concerns": r.major_concerns,
        "minor_concerns": r.minor_concerns, "recommendation": r.recommendation,
        "score_novelty": r.score_novelty, "score_methodology": r.score_methodology,
        "score_clarity": r.score_clarity, "score_statistics": r.score_statistics,
        "claude_analysis": r.claude_analysis, "claude_checklist": r.claude_checklist,
        "status": r.status, "updated_at": r.updated_at.isoformat(),
        "due_date": r.paper.due_date.isoformat() if r.paper and r.paper.due_date else None
    }
